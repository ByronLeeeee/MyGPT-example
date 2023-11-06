from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.docstore.document import Document
from langchain.prompts import PromptTemplate
from langchain.chains.question_answering import load_qa_chain
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA


import os
import docx

from langchain.document_loaders import TextLoader


#如需设置网络代理，请删掉前方的#号并设置具体的代理信息，如果没有网络问题可以不用修改；
# os.environ["http_proxy"] = "http://127.0.0.1:xxxx"
# os.environ["https_proxy"] = "http://127.0.0.1:xxxx"

os.environ["OPENAI_API_KEY"] = "YOUR-API-KEY" #设置申领到的OpenAI API Key

class LawTextLoader(TextLoader):
    def load(self):
        with open(self.file_path, encoding=self.encoding) as f: #打开txt文本
            text = f.readlines() #逐行读取文字并保存到text列表中
        metadata = {"source": self.file_path} #创建metadata字典用于保存文本来源
        doc = []
        for t in text:
            #把text列表中保存的每一条法条以Document格式保存到doc列表中
            doc.append(Document(page_content = t, metadata = metadata)) 
        return doc

LAWS_PATH = './laws' #设定放置法条的文件夹的路径，这里是相对路径

def docx2txt(laws_path):
    #新建一个用于记录文档名称的files列表，里面每个值都来源于传进去的文件夹路径中，后缀是.docx的文件
    files = [file for file in os.listdir(laws_path) if file.endswith('.docx')]

    for file in files: #遍历files列表中的每一个值
        docx_file_path = os.path.join(laws_path, file) #把路径和文件名称拼合
        doc = docx.Document(docx_file_path) #读取docx文档的内容
        law_name = "《" + file.replace('.docx', '')  + "》"  # 获取法律名称
        print("正在处理%s" % law_name) #方便在终端中观察正在处理的文件名称
        last_para = None  # 重置用于储存段落信息的变量
        text_total = []  # 用于存储完整的文本

        for para in doc.paragraphs: #遍历docx文档中的全部段落
            #如果有前一段的数据，且该段的内容不是以‘第’字开头的话，则last_para变量中增加这一段的内容，并进入下一个遍历
            if last_para and not (para.text.startswith('第') or para.text.startswith('　　第')):
                last_para += para.text
                continue

            #如果有前一段的数据，且该段的内容以‘第’字开头的话，则在text_total的列表中增加本法律的名称以及上一段的内容，并以换行结尾
            if last_para and (para.text.startswith('第') or para.text.startswith('　　第')):
                text_total.append(f'{law_name} {last_para}\n')  # 添加法律名称和段落文本
            
            #都不符合上面的两个情况，则把这段的内容保存到last_para变量中
            last_para = para.text
        text_total.append(f'{law_name} {last_para}')  # 添加最后一个段落文本

        txt_file_path = os.path.join(laws_path, f'{law_name}.txt')  # 构建输出文本文件路径
        with open(txt_file_path, 'w', encoding='utf-8') as f: #创建一个文本文件
            f.write(''.join(text_total))  # 写入文本内容

def lawEmbedding(laws_path):
    # 引入进度条库，方便我们观察进度
    from tqdm import tqdm
    
    # 新建一个用于记录文档名称的files列表，里面每个值都来源于传进去的文件夹路径中，后缀是.txt的文件
    files = [file for file in os.listdir(laws_path) if file.endswith('.txt')]
    total_files = len(files) # 计算总文件个数
    
    # 对于每个文件，使用进度条进行处理
    for i, file in tqdm(enumerate(files, start=1), total=total_files, desc='向量存储'):
        print('正在处理 (%d/%d)' % (i, total_files))
        print("发现文件: " + file)
        file_path = os.path.join(laws_path, file)
        
        # 使用我们继承的LawTextLoader读取并格式化txt文件
        loader = LawTextLoader(file_path, encoding='utf-8')
        docs = loader.load()
        
        # 设定向量数据库的名称为Law_DB
        persist_directory = 'Law_DB'

        # 使用不同的Embedding模型将会有不同的效果
        embedding = OpenAIEmbeddings()
        # 以下为比较优秀的开源Embedding模型，更专注于中文
        # from langchain.embeddings import HuggingFaceEmbeddings
        # embedding = HuggingFaceEmbeddings(model_name="GanymedeNil/text2vec-large-chinese") 
        
        # 把法条存入数据库
        vectordb = Chroma.from_documents(
            documents=docs, embedding=embedding, persist_directory=persist_directory)
        
        vectordb = None  # 清除向量数据库的引用，释放内存
        
        print(file + " 向量完成")
    
    # 所有法条向量保存完成
    print("全部法条向量保存完成！")
    
    # 输出保存的文件列表
    result = '本次成功完成以下文件的录入：\n'+'\n'.join(files)
    return result

def LawQA(query):
    # 为ChatGPT设置模版Prompt
    PROMPT = PromptTemplate(
        input_variables=["context", "question"],
        template= """你是一名律师助理，请分析Question所蕴含的法律事实并参考Context提供的法律条文进行回答"。
                    Context:{context},
                    Question: {question}
                    请用中文按以下格式回复：
                    法律依据:
                    分析： """
                    )

    # 选择Embedding模型
    embedding = OpenAIEmbeddings()
    # embedding = HuggingFaceEmbeddings(model_name="GanymedeNil/text2vec-large-chinese")
    # 载入数据库
    vectordb = Chroma(persist_directory='Law_DB', embedding_function=embedding)
    # k值为让数据库拿出来相似度最高的参考材料的数量，不宜太多，不然占用token数会太多
    retriever = vectordb.as_retriever(search_kwargs={"k": 3})

    # 需要通过from langchain.chains import RetrievalQA引入RetrievalQA函数
    # 创建一个问答的链，并指定使用ChatGPT作为模型
    qa = RetrievalQA.from_chain_type(llm= ChatOpenAI(temperature=0.1)
        # temperature:GPT的回复创作度，0-1，越低越严谨；max_tokens：最大接受的tokens;
        # chain_type:链的回答模式，stuff-直接分析，map_reduce - 多段分析，refine - 归纳
        , retriever=retriever, chain_type="stuff", chain_type_kwargs={"prompt": PROMPT})
    # 开始让问答链回答我们的问题
    result = qa.run(query)
    return result

if __name__ == "__main__":
    # 循环输入查询，直到输入 "退出"
    print("您好，我是您的AI法律助手，请输入问题： (或输入 '退出' 结束): \n\n")
    while True:
        query = input("问题:")
        if query == '退出':
            print('再见！')
            break
        print('\n回答:\n' + LawQA(query) + '\n')
